# =========================================================================
#  Библиотека — Чтение и поиск / Library — Reading and Search
# =========================================================================
import sys
import os
import re
import json
import uuid

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLineEdit, QListWidget, QListWidgetItem, QLabel,
    QComboBox, QSplitter, QTextBrowser, QFileDialog, QMessageBox,
    QInputDialog, QDialog
)
from PyQt6.QtCore import Qt


# =========================================================================
#  ПЕРЕВОДЫ / TRANSLATIONS
# =========================================================================
TRANSLATIONS = {
    'ru': {
        'window_title':         "Библиотека — Чтение и поиск",
        'search_placeholder':   "Слово или его часть (напр. Катари)…",
        'btn_add_books':        "＋ Добавить книги",
        'btn_search':           "Искать",
        'btn_about':            "ℹ О программе",
        'lbl_books':            "Открытые книги",
        'btn_remove_book':      "－ Убрать книгу",
        'lbl_bookmarks':        "Подборки-закладки",
        'btn_new_bm':           "Создать подборку",
        'btn_rename_bm':        "✎ Переименовать подборку",
        'btn_delete_bm':        "✖ Удалить подборку",
        'btn_export_bm':        "⤓ Экспорт подборки",
        'lbl_results':          "Результаты поиска",
        'btn_prev':             "◀ Назад",
        'btn_next':             "Вперёд ▶",
        'btn_add_to_bm':        "★ Добавить в подборку",
        'btn_remove_from_bm':   "🗑 Удалить из подборки",
        'lbl_reading':          "Чтение",
        'btn_expand_up':        "◀ Абзац сверху",
        'btn_shrink_up':        "Убрать сверху",
        'btn_shrink_down':      "Убрать снизу",
        'btn_expand_down':      "Абзац снизу ▶",
        'page_label':           "Страница {cur} из {total}",
        'lbl_results_search':   "Результаты поиска — найдено: {n}",
        'lbl_results_bookmark': "Подборка — найдено: {n}",
        'lbl_reading_title':    "Чтение — {title}",
        # О программе
        'about_title':          "О программе",
        'about_app':            "Библиотека — Чтение и поиск",
        'about_version':        "Версия 1.0",
        'about_studio':         "Создано студией  <b>NeuroOctopus</b>",
        # Диалоги
        'dlg_open_fail':        "Не удалось открыть",
        'dlg_bm_info':          "Подборка",
        'dlg_new_bm_title':     "Новая подборка",
        'dlg_new_bm_prompt':    "Название подборки:",
        'dlg_rename_title':     "Переименовать подборку",
        'dlg_rename_prompt':    "Новое название:",
        'dlg_rename_done':      "Подборка переименована в «{name}».",
        'dlg_rename_exists':    "Подборка «{name}» уже существует.",
        'dlg_delete_title':     "Удалить подборку",
        'dlg_delete_q':         "Удалить подборку «{name}»{count} целиком?",
        'dlg_delete_count':     " ({n} фрагм.)",
        'dlg_delete_done':      "Подборка «{name}» удалена.",
        'dlg_no_bm':            "Нет выбранной подборки.",
        'dlg_select_frag':      "Сначала выберите фрагмент.",
        'dlg_create_bm':        "Сначала создайте подборку.",
        'dlg_already_in_bm':    "Этот фрагмент уже есть в подборке «{name}».",
        'dlg_added_to_bm':      "Добавлено в подборку «{name}».",
        'dlg_removed_from_bm':  "Удалено из подборки «{name}».",
        'dlg_not_in_bm':        "Этого фрагмента нет в выбранной подборке.",
        'dlg_export_title':     "Экспорт",
        'dlg_export_empty':     "Подборка пуста.",
        'dlg_export_save':      "Экспорт подборки",
        'dlg_export_filters':   "Текст (*.txt);;Word (*.docx);;PDF (*.pdf)",
        'dlg_export_done':      "Готово.",
        'dlg_export_err':       "Ошибка: {e}",
        'export_heading':       "Подборка: {name}",
        'file_filter_books':    ("Книги (*.txt *.epub *.fb2 *.pdf *.docx "
                                 "*.djvu *.lrf *.azw *.azw3 *.mobi);;"
                                 "Все файлы (*)"),
        'file_open_title':      "Выберите книги",
    },
    'en': {
        'window_title':         "Library — Reading and Search",
        'search_placeholder':   "Word or part of it (e.g. Cathar)…",
        'btn_add_books':        "＋ Add books",
        'btn_search':           "Search",
        'btn_about':            "ℹ About",
        'lbl_books':            "Opened books",
        'btn_remove_book':      "－ Remove book",
        'lbl_bookmarks':        "Collections",
        'btn_new_bm':           "New collection",
        'btn_rename_bm':        "✎ Rename collection",
        'btn_delete_bm':        "✖ Delete collection",
        'btn_export_bm':        "⤓ Export collection",
        'lbl_results':          "Search results",
        'btn_prev':             "◀ Back",
        'btn_next':             "Forward ▶",
        'btn_add_to_bm':        "★ Add to collection",
        'btn_remove_from_bm':   "🗑 Remove from collection",
        'lbl_reading':          "Reading",
        'btn_expand_up':        "◀ Paragraph above",
        'btn_shrink_up':        "Remove above",
        'btn_shrink_down':      "Remove below",
        'btn_expand_down':      "Paragraph below ▶",
        'page_label':           "Page {cur} of {total}",
        'lbl_results_search':   "Search results — found: {n}",
        'lbl_results_bookmark': "Collection — found: {n}",
        'lbl_reading_title':    "Reading — {title}",
        # About
        'about_title':          "About",
        'about_app':            "Library — Reading and Search",
        'about_version':        "Version 1.0",
        'about_studio':         "Created by  <b>NeuroOctopus</b>",
        # Dialogs
        'dlg_open_fail':        "Could not open",
        'dlg_bm_info':          "Collection",
        'dlg_new_bm_title':     "New collection",
        'dlg_new_bm_prompt':    "Collection name:",
        'dlg_rename_title':     "Rename collection",
        'dlg_rename_prompt':    "New name:",
        'dlg_rename_done':      "Collection renamed to «{name}».",
        'dlg_rename_exists':    "Collection «{name}» already exists.",
        'dlg_delete_title':     "Delete collection",
        'dlg_delete_q':         "Delete collection «{name}»{count} entirely?",
        'dlg_delete_count':     " ({n} fragment(s))",
        'dlg_delete_done':      "Collection «{name}» deleted.",
        'dlg_no_bm':            "No collection selected.",
        'dlg_select_frag':      "Please select a fragment first.",
        'dlg_create_bm':        "Please create a collection first.",
        'dlg_already_in_bm':    "This fragment is already in «{name}».",
        'dlg_added_to_bm':      "Added to collection «{name}».",
        'dlg_removed_from_bm':  "Removed from collection «{name}».",
        'dlg_not_in_bm':        "This fragment is not in the selected collection.",
        'dlg_export_title':     "Export",
        'dlg_export_empty':     "Collection is empty.",
        'dlg_export_save':      "Export collection",
        'dlg_export_filters':   "Text (*.txt);;Word (*.docx);;PDF (*.pdf)",
        'dlg_export_done':      "Done.",
        'dlg_export_err':       "Error: {e}",
        'export_heading':       "Collection: {name}",
        'file_filter_books':    ("Books (*.txt *.epub *.fb2 *.pdf *.docx "
                                 "*.djvu *.lrf *.azw *.azw3 *.mobi);;"
                                 "All files (*)"),
        'file_open_title':      "Select books",
    },
}


def tr(lang, key, **kwargs):
    """Получить строку перевода."""
    text = TRANSLATIONS.get(lang, TRANSLATIONS['ru']).get(key, key)
    return text.format(**kwargs) if kwargs else text


# =========================================================================
#  ЗАГРУЗКА КНИГ
# =========================================================================
class BookLoader:
    @staticmethod
    def load(path):
        ext = os.path.splitext(path)[1].lower()
        if ext == ".txt":  return BookLoader._txt(path)
        if ext == ".epub": return BookLoader._epub(path)
        if ext == ".fb2":  return BookLoader._fb2(path)
        if ext == ".pdf":  return BookLoader._pdf(path)
        if ext == ".docx": return BookLoader._docx(path)
        if ext == ".djvu": return BookLoader._djvu(path)
        if ext in (".azw", ".azw3", ".mobi", ".lrf"):
            return BookLoader._ebook(path)
        raise ValueError(f"Формат {ext} не поддерживается")

    @staticmethod
    def _split(text):
        paras = re.split(r'\n\s*\n', text.replace('\r\n', '\n'))
        return [p.strip() for p in paras if p.strip()]

    @staticmethod
    def _txt(path):
        for enc in ("utf-8", "cp1251", "latin-1"):
            try:
                with open(path, encoding=enc) as f:
                    return BookLoader._split(f.read())
            except Exception:
                continue
        raise IOError("Не удалось прочитать TXT")

    @staticmethod
    def _epub(path):
        try:
            from ebooklib import epub
            import ebooklib
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("pip install EbookLib beautifulsoup4")
        book = epub.read_epub(path)
        text = []
        for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
            soup = BeautifulSoup(item.get_content(), "html.parser")
            text.append(soup.get_text("\n"))
        return BookLoader._split("\n\n".join(text))

    @staticmethod
    def _fb2(path):
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("pip install beautifulsoup4 lxml")
        with open(path, encoding="utf-8", errors="ignore") as f:
            soup = BeautifulSoup(f.read(), "lxml-xml")
        paras = [p.get_text(" ", strip=True) for p in soup.find_all("p")]
        return [p for p in paras if p]

    @staticmethod
    def _pdf(path):
        try:
            import fitz
        except ImportError:
            raise ImportError("pip install PyMuPDF")
        doc = fitz.open(path)
        text = "\n\n".join(page.get_text() for page in doc)
        doc.close()
        return BookLoader._split(text)

    @staticmethod
    def _docx(path):
        try:
            import docx
        except ImportError:
            raise ImportError("pip install python-docx")
        doc = docx.Document(path)
        return [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    @staticmethod
    def _djvu(path):
        try:
            import subprocess
            out = subprocess.check_output(["djvutxt", path])
            return BookLoader._split(out.decode("utf-8", errors="ignore"))
        except Exception:
            raise ImportError("Нужна утилита djvutxt (DjVuLibre)")

    @staticmethod
    def _ebook(path):
        try:
            import subprocess
            import tempfile
            tmp = tempfile.mktemp(suffix=".txt")
            subprocess.check_call(["ebook-convert", path, tmp])
            with open(tmp, encoding="utf-8", errors="ignore") as f:
                data = f.read()
            os.remove(tmp)
            return BookLoader._split(data)
        except Exception:
            raise ImportError("Нужен ebook-convert (Calibre)")


# =========================================================================
#  КНИГА
# =========================================================================
class Book:
    def __init__(self, path):
        self.path = path
        self.title = os.path.splitext(os.path.basename(path))[0]
        self.paragraphs = BookLoader.load(path)
        if not self.paragraphs:
            raise ValueError("Книга пуста или текст не извлечён")


# =========================================================================
#  СТИЛЬ
# =========================================================================
STYLE = """
    QWidget { background:#F5F1E8; color:#2B2B2B;
              font-family:'Segoe UI',sans-serif; font-size:14px; }
    QDialog { background:#F5F1E8; }
    QPushButton { background:#A65A42; color:#F5F1E8; border:none;
                  padding:8px 14px; border-radius:4px; }
    QPushButton:hover { background:#8F4A34; }
    QPushButton:disabled { background:#CBBFAE; color:#F5F1E8; }
    QLineEdit, QComboBox { background:#FFF; border:1px solid #D6CBB5;
                  padding:8px; border-radius:4px; }
    QListWidget { background:#FBF8F1; border:1px solid #D6CBB5;
                  border-radius:4px; }
    QListWidget::item { padding:8px; border-bottom:1px solid #EDE5D3; }
    QListWidget::item:selected { background:#EFE3D0; color:#2B2B2B; }
    QLabel { font-weight:bold; padding:4px 0; }
    QTextBrowser { background:#FBF8F1; border:1px solid #D6CBB5; padding:15px; }
"""


# =========================================================================
#  Диалог "О программе" / About
# =========================================================================
class AboutDialog(QDialog):
    def __init__(self, lang='ru', parent=None):
        super().__init__(parent)
        self.setWindowTitle(tr(lang, 'about_title'))
        self.setFixedSize(380, 220)
        self.setStyleSheet(STYLE)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(30, 25, 30, 20)
        layout.setSpacing(10)

        title = QLabel(tr(lang, 'about_app'))
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        title.setStyleSheet("font-size:15px; font-weight:bold; color:#A65A42;")
        layout.addWidget(title)

        version = QLabel(tr(lang, 'about_version'))
        version.setAlignment(Qt.AlignmentFlag.AlignCenter)
        version.setStyleSheet("font-size:12px; font-weight:normal;")
        layout.addWidget(version)

        line = QLabel()
        line.setFixedHeight(1)
        line.setStyleSheet("background:#D6CBB5; margin:4px 0;")
        layout.addWidget(line)

        studio = QLabel(tr(lang, 'about_studio'))
        studio.setAlignment(Qt.AlignmentFlag.AlignCenter)
        studio.setStyleSheet("font-size:13px; font-weight:normal;")
        layout.addWidget(studio)

        tg = QLabel('<a href="https://t.me/neuroctopus" '
                    'style="color:#A65A42;">t.me/neuroctopus</a>')
        tg.setAlignment(Qt.AlignmentFlag.AlignCenter)
        tg.setOpenExternalLinks(True)
        tg.setStyleSheet("font-size:13px;")
        layout.addWidget(tg)

        layout.addStretch()

        btn_ok = QPushButton("OK")
        btn_ok.setFixedWidth(90)
        btn_ok.clicked.connect(self.accept)
        row = QHBoxLayout()
        row.addStretch()
        row.addWidget(btn_ok)
        row.addStretch()
        layout.addLayout(row)


# =========================================================================
#  ГЛАВНОЕ ОКНО
# =========================================================================
class Reader(QMainWindow):
    RESULTS_PER_PAGE = 30
    SESSION_FILE = "session.json"

    def __init__(self):
        super().__init__()
        self.lang = 'ru'          # текущий язык
        self.books = []
        self.search_results = []
        self.current_page = 0
        self.current_bookmark = None
        self.bookmarks = {}
        self.reader_font_size = 13
        self._last_fragment = None
        self._cur_start = 0
        self._cur_end = 0

        self.resize(1300, 850)
        self.setStyleSheet(STYLE)

        # Иконка
        from PyQt6.QtGui import QIcon
        if getattr(sys, 'frozen', False):
            base = sys._MEIPASS
        else:
            base = os.path.dirname(os.path.abspath(__file__))
        icon_path = os.path.join(base, "icon.ico")
        if os.path.exists(icon_path):
            self.setWindowIcon(QIcon(icon_path))

        self.init_ui()
        self.load_session()
        self._apply_lang()   # применяем язык после загрузки (он мог быть сохранён)

    # ------------------------------------------------------------------
    #  Утилита перевода
    # ------------------------------------------------------------------
    def t(self, key, **kwargs):
        return tr(self.lang, key, **kwargs)

    # ------------------------------------------------------------------
    #  Переключение языка
    # ------------------------------------------------------------------
    def _on_lang_changed(self, text):
        self.lang = 'en' if text == 'EN' else 'ru'
        self._apply_lang()
        self.save_session()

    def _apply_lang(self):
        """Обновить все надписи под текущий язык."""
        # Синхронизировать комбо без лишнего сигнала
        self.lang_combo.blockSignals(True)
        self.lang_combo.setCurrentText('EN' if self.lang == 'en' else 'RU')
        self.lang_combo.blockSignals(False)

        self.setWindowTitle(self.t('window_title'))
        self.search_field.setPlaceholderText(self.t('search_placeholder'))

        # Верхняя панель
        self.btn_add_books.setText(self.t('btn_add_books'))
        self.btn_search.setText(self.t('btn_search'))
        self.btn_about.setText(self.t('btn_about'))

        # Левая колонка
        self.lbl_books.setText(self.t('lbl_books'))
        self.btn_remove_book.setText(self.t('btn_remove_book'))
        self.lbl_bookmarks.setText(self.t('lbl_bookmarks'))
        self.btn_new_bm.setText(self.t('btn_new_bm'))
        self.btn_rename_bm.setText(self.t('btn_rename_bm'))
        self.btn_delete_bm.setText(self.t('btn_delete_bm'))
        self.btn_export_bm.setText(self.t('btn_export_bm'))

        # Центр
        self.btn_prev.setText(self.t('btn_prev'))
        self.btn_next.setText(self.t('btn_next'))
        self.btn_add_to_bm.setText(self.t('btn_add_to_bm'))
        self.btn_remove_from_bm.setText(self.t('btn_remove_from_bm'))

        # Правая панель — expand-кнопки
        self.up_btn.setText(self.t('btn_expand_up'))
        self.up_rem_btn.setText(self.t('btn_shrink_up'))
        self.down_rem_btn.setText(self.t('btn_shrink_down'))
        self.down_btn.setText(self.t('btn_expand_down'))

        # Заголовок читалки
        if self._last_fragment:
            self.reader_label.setText(
                self.t('lbl_reading_title',
                       title=self._last_fragment['book'].title))
        else:
            self.reader_label.setText(self.t('lbl_reading'))

        # Перерисовать список результатов (метки страниц тоже переведутся)
        self.render_results()

    # ------------------------------------------------------------------
    #  Построение интерфейса
    # ------------------------------------------------------------------
    def _btn(self, text, slot, width=None):
        b = QPushButton(text)
        b.clicked.connect(slot)
        if width:
            b.setFixedWidth(width)
        return b

    def init_ui(self):
        central = QWidget()
        self.setCentralWidget(central)
        main = QVBoxLayout(central)
        main.setContentsMargins(15, 15, 15, 15)
        main.setSpacing(10)

        # ── Верхняя панель ──────────────────────────────────────────────
        top = QHBoxLayout()

        self.btn_add_books = QPushButton()
        self.btn_add_books.clicked.connect(self.add_books)

        self.search_field = QLineEdit()
        self.search_field.returnPressed.connect(self.do_search)

        self.btn_search = QPushButton()
        self.btn_search.clicked.connect(self.do_search)

        self.btn_about = QPushButton()
        self.btn_about.clicked.connect(self.show_about)

        # Переключатель языка
        self.lang_combo = QComboBox()
        self.lang_combo.addItems(["RU", "EN"])
        self.lang_combo.setFixedWidth(68)
        self.lang_combo.currentTextChanged.connect(self._on_lang_changed)

        top.addWidget(self.btn_add_books)
        top.addWidget(self.search_field, 1)
        top.addWidget(self.btn_search)
        top.addWidget(self.btn_about)
        top.addWidget(self.lang_combo)
        main.addLayout(top)

        splitter = QSplitter(Qt.Orientation.Horizontal)

        # ── Левая колонка ───────────────────────────────────────────────
        left = QWidget()
        ll = QVBoxLayout(left)

        self.lbl_books = QLabel()
        ll.addWidget(self.lbl_books)

        self.book_list = QListWidget()
        ll.addWidget(self.book_list)

        self.btn_remove_book = QPushButton()
        self.btn_remove_book.clicked.connect(self.remove_book)
        ll.addWidget(self.btn_remove_book)

        self.lbl_bookmarks = QLabel()
        ll.addWidget(self.lbl_bookmarks)

        self.bm_combo = QComboBox()
        self.bm_combo.currentTextChanged.connect(self.show_bookmark)
        ll.addWidget(self.bm_combo)

        self.btn_new_bm = QPushButton()
        self.btn_new_bm.clicked.connect(self.new_bookmark)
        ll.addWidget(self.btn_new_bm)

        self.btn_rename_bm = QPushButton()
        self.btn_rename_bm.clicked.connect(self.rename_bookmark)
        ll.addWidget(self.btn_rename_bm)

        self.btn_delete_bm = QPushButton()
        self.btn_delete_bm.clicked.connect(self.delete_bookmark)
        ll.addWidget(self.btn_delete_bm)

        self.btn_export_bm = QPushButton()
        self.btn_export_bm.clicked.connect(self.export_bookmark)
        ll.addWidget(self.btn_export_bm)

        splitter.addWidget(left)

        # ── Центр — результаты ──────────────────────────────────────────
        center = QWidget()
        cl = QVBoxLayout(center)

        self.results_label = QLabel()
        cl.addWidget(self.results_label)

        self.results_list = QListWidget()
        self.results_list.setWordWrap(True)
        self.results_list.itemDoubleClicked.connect(self.open_fragment)
        cl.addWidget(self.results_list)

        nav = QHBoxLayout()
        self.btn_prev = QPushButton()
        self.btn_prev.clicked.connect(self.prev_page)
        self.page_label = QLabel("—")
        self.btn_next = QPushButton()
        self.btn_next.clicked.connect(self.next_page)
        nav.addWidget(self.btn_prev)
        nav.addWidget(self.page_label, 1, Qt.AlignmentFlag.AlignCenter)
        nav.addWidget(self.btn_next)
        cl.addLayout(nav)

        self.btn_add_to_bm = QPushButton()
        self.btn_add_to_bm.clicked.connect(self.add_to_bookmark)
        cl.addWidget(self.btn_add_to_bm)

        self.btn_remove_from_bm = QPushButton()
        self.btn_remove_from_bm.clicked.connect(self.remove_from_bookmark)
        cl.addWidget(self.btn_remove_from_bm)

        splitter.addWidget(center)

        # ── Правая — чтение ─────────────────────────────────────────────
        right = QWidget()
        rl = QVBoxLayout(right)

        head = QHBoxLayout()
        self.reader_label = QLabel()
        head.addWidget(self.reader_label, 1)
        head.addWidget(self._btn("А−", self.decrease_font, 50))
        head.addWidget(self._btn("А+", self.increase_font, 50))
        rl.addLayout(head)

        exp = QHBoxLayout()
        self.up_btn       = QPushButton()
        self.up_btn.clicked.connect(self.expand_up)
        self.up_rem_btn   = QPushButton()
        self.up_rem_btn.clicked.connect(self.shrink_up)
        self.down_rem_btn = QPushButton()
        self.down_rem_btn.clicked.connect(self.shrink_down)
        self.down_btn     = QPushButton()
        self.down_btn.clicked.connect(self.expand_down)
        exp.addWidget(self.up_btn)
        exp.addWidget(self.up_rem_btn)
        exp.addWidget(self.down_rem_btn)
        exp.addWidget(self.down_btn)
        rl.addLayout(exp)

        self.text_view = QTextBrowser()
        rl.addWidget(self.text_view)
        splitter.addWidget(right)

        splitter.setSizes([260, 430, 610])
        main.addWidget(splitter, 1)
        self._update_expand_buttons()

    # ------------------------------------------------------------------
    def show_about(self):
        dlg = AboutDialog(lang=self.lang, parent=self)
        dlg.exec()

    # ------------------------------------------------------------------
    #  Книги
    # ------------------------------------------------------------------
    def add_books(self):
        paths, _ = QFileDialog.getOpenFileNames(
            self,
            self.t('file_open_title'), "",
            self.t('file_filter_books'))
        for p in paths:
            try:
                b = Book(p)
                self.books.append(b)
                self.book_list.addItem(b.title)
            except Exception as e:
                QMessageBox.warning(self, self.t('dlg_open_fail'),
                                    f"{os.path.basename(p)}:\n{e}")
        self.save_session()

    def remove_book(self):
        row = self.book_list.currentRow()
        if row >= 0:
            self.book_list.takeItem(row)
            del self.books[row]
            self.save_session()

        # ------------------------------------------------------------------
    #  Поиск
    # ------------------------------------------------------------------
    def do_search(self):
        self.current_bookmark = None
        query = self.search_field.text().strip()
        if not query:
            return
        pat = re.compile(r'\b' + re.escape(query) + r'\w*', re.IGNORECASE)
        self.search_results = [
            {'book': b, 'index': i, 'text': para,
             'query': query, 'start': i, 'end': i}
            for b in self.books
            for i, para in enumerate(b.paragraphs)
            if pat.search(para)]
        self.current_page = 0
        self.render_results()

    # ------------------------------------------------------------------
    #  Отображение результатов
    # ------------------------------------------------------------------
    def render_results(self):
        self.results_list.clear()
        total = len(self.search_results)
        pages = max(1, (total + self.RESULTS_PER_PAGE - 1)
                    // self.RESULTS_PER_PAGE)
        self.current_page = max(0, min(self.current_page, pages - 1))

        start = self.current_page * self.RESULTS_PER_PAGE
        end   = start + self.RESULTS_PER_PAGE

        if self.current_bookmark is None:
            self.results_label.setText(
                self.t('lbl_results_search', n=total))
        else:
            self.results_label.setText(
                self.t('lbl_results_bookmark', n=total))

        for d in self.search_results[start:end]:
            preview = d['text'][:160].replace('\n', ' ')
            label   = f"[{d['book'].title}]  {preview}…"
            item    = QListWidgetItem(label)
            item.setData(Qt.ItemDataRole.UserRole, d)
            self.results_list.addItem(item)

        if total > 0:
            self.page_label.setText(
                self.t('page_label',
                       cur=self.current_page + 1,
                       total=pages))
        else:
            self.page_label.setText("—")

        self.btn_prev.setEnabled(self.current_page > 0)
        self.btn_next.setEnabled(self.current_page < pages - 1)

    def prev_page(self):
        if self.current_page > 0:
            self.current_page -= 1
            self.render_results()

    def next_page(self):
        total = len(self.search_results)
        pages = max(1, (total + self.RESULTS_PER_PAGE - 1)
                    // self.RESULTS_PER_PAGE)
        if self.current_page < pages - 1:
            self.current_page += 1
            self.render_results()

    # ------------------------------------------------------------------
    #  Открытие фрагмента в читалке
    # ------------------------------------------------------------------
    def open_fragment(self, item):
        d = item.data(Qt.ItemDataRole.UserRole)
        self._last_fragment = d
        self._cur_start = d.get('start', d['index'])
        self._cur_end   = d.get('end',   d['index'])
        self.reader_label.setText(
            self.t('lbl_reading_title', title=d['book'].title))
        self._render_reader()
        self._update_expand_buttons()

    def _render_reader(self):
        if self._last_fragment is None:
            return
        book  = self._last_fragment['book']
        query = self._last_fragment.get('query', '')
        paras = book.paragraphs[self._cur_start: self._cur_end + 1]
        html_parts = []
        for para in paras:
            escaped = (para
                       .replace('&', '&amp;')
                       .replace('<', '&lt;')
                       .replace('>', '&gt;'))
            if query:
                pat = re.compile(
                    r'(' + re.escape(query) + r'\w*)', re.IGNORECASE)
                escaped = pat.sub(
                    r'<mark style="background:#F5C518;'
                    r'border-radius:3px;">\1</mark>',
                    escaped)
            html_parts.append(f"<p>{escaped}</p>")

        html = (f'<div style="font-size:{self.reader_font_size}pt;'
                f'line-height:1.7;">'
                + ''.join(html_parts) + '</div>')
        self.text_view.setHtml(html)

    # ------------------------------------------------------------------
    #  Expand / Shrink абзацев
    # ------------------------------------------------------------------
    def _update_expand_buttons(self):
        has = self._last_fragment is not None
        book = self._last_fragment['book'] if has else None

        can_up   = has and self._cur_start > 0
        can_down = has and self._cur_end < len(book.paragraphs) - 1
        can_sh_u = has and self._cur_start < self._cur_end
        can_sh_d = has and self._cur_start < self._cur_end

        self.up_btn.setEnabled(can_up)
        self.down_btn.setEnabled(can_down)
        self.up_rem_btn.setEnabled(can_sh_u)
        self.down_rem_btn.setEnabled(can_sh_d)

    def expand_up(self):
        if self._last_fragment and self._cur_start > 0:
            self._cur_start -= 1
            self._render_reader()
            self._update_expand_buttons()

    def expand_down(self):
        if self._last_fragment:
            book = self._last_fragment['book']
            if self._cur_end < len(book.paragraphs) - 1:
                self._cur_end += 1
                self._render_reader()
                self._update_expand_buttons()

    def shrink_up(self):
        if self._last_fragment and self._cur_start < self._cur_end:
            self._cur_start += 1
            self._render_reader()
            self._update_expand_buttons()

    def shrink_down(self):
        if self._last_fragment and self._cur_start < self._cur_end:
            self._cur_end -= 1
            self._render_reader()
            self._update_expand_buttons()

    # ------------------------------------------------------------------
    #  Размер шрифта
    # ------------------------------------------------------------------
    def increase_font(self):
        self.reader_font_size = min(self.reader_font_size + 1, 32)
        self._render_reader()

    def decrease_font(self):
        self.reader_font_size = max(self.reader_font_size - 1, 8)
        self._render_reader()

    # ------------------------------------------------------------------
    #  Подборки / Bookmarks
    # ------------------------------------------------------------------
    def new_bookmark(self):
        name, ok = QInputDialog.getText(
            self, self.t('dlg_new_bm_title'), self.t('dlg_new_bm_prompt'))
        if ok and name.strip():
            name = name.strip()
            if name not in self.bookmarks:
                self.bookmarks[name] = []
                self.bm_combo.addItem(name)
                self.bm_combo.setCurrentText(name)
            self.save_session()

    def rename_bookmark(self):
        old = self.bm_combo.currentText()
        if not old:
            QMessageBox.information(self, self.t('dlg_bm_info'),
                                    self.t('dlg_no_bm'))
            return
        new, ok = QInputDialog.getText(
            self, self.t('dlg_rename_title'),
            self.t('dlg_rename_prompt'), text=old)
        if not ok or not new.strip():
            return
        new = new.strip()
        if new == old:
            return
        if new in self.bookmarks:
            QMessageBox.warning(self, self.t('dlg_bm_info'),
                                self.t('dlg_rename_exists', name=new))
            return
        self.bookmarks[new] = self.bookmarks.pop(old)
        idx = self.bm_combo.findText(old)
        self.bm_combo.setItemText(idx, new)
        self.bm_combo.setCurrentText(new)
        QMessageBox.information(self, self.t('dlg_bm_info'),
                                self.t('dlg_rename_done', name=new))
        self.save_session()

    def delete_bookmark(self):
        name = self.bm_combo.currentText()
        if not name:
            QMessageBox.information(self, self.t('dlg_bm_info'),
                                    self.t('dlg_no_bm'))
            return
        n     = len(self.bookmarks.get(name, []))
        count = self.t('dlg_delete_count', n=n) if n else ""
        reply = QMessageBox.question(
            self,
            self.t('dlg_delete_title'),
            self.t('dlg_delete_q', name=name, count=count))
        if reply == QMessageBox.StandardButton.Yes:
            del self.bookmarks[name]
            idx = self.bm_combo.findText(name)
            self.bm_combo.removeItem(idx)
            self.search_results = []
            self.render_results()
            QMessageBox.information(self, self.t('dlg_bm_info'),
                                    self.t('dlg_delete_done', name=name))
            self.save_session()

    def show_bookmark(self, name):
        if not name or name not in self.bookmarks:
            return
        self.current_bookmark = name
        self.search_results = list(self.bookmarks[name])
        self.current_page   = 0
        self.render_results()

    def _fragment_key(self, d):
        """Уникальный ключ фрагмента для сравнения."""
        return (d['book'].path,
                d.get('orig_index', d.get('index')))

    def add_to_bookmark(self):
        name = self.bm_combo.currentText()
        if not name:
            QMessageBox.information(self, self.t('dlg_bm_info'),
                                    self.t('dlg_create_bm'))
            return
        item = self.results_list.currentItem()
        if item is None:
            QMessageBox.information(self, self.t('dlg_bm_info'),
                                    self.t('dlg_select_frag'))
            return
        d   = item.data(Qt.ItemDataRole.UserRole)
        key = self._fragment_key(d)
        existing_keys = [self._fragment_key(s)
                         for s in self.bookmarks[name]]
        if key in existing_keys:
            QMessageBox.information(self, self.t('dlg_bm_info'),
                                    self.t('dlg_already_in_bm', name=name))
            return
        # Сохранить путь книги для восстановления из сессии
        d['path']       = d['book'].path
        d['orig_index'] = d.get('index')
        self.bookmarks[name].append(d)
        QMessageBox.information(self, self.t('dlg_bm_info'),
                                self.t('dlg_added_to_bm', name=name))
        self.save_session()

    def remove_from_bookmark(self):
        name = self.bm_combo.currentText()
        if not name:
            QMessageBox.information(self, self.t('dlg_bm_info'),
                                    self.t('dlg_no_bm'))
            return
        item = self.results_list.currentItem()
        if item is None:
            QMessageBox.information(self, self.t('dlg_bm_info'),
                                    self.t('dlg_select_frag'))
            return
        d   = item.data(Qt.ItemDataRole.UserRole)
        key = self._fragment_key(d)
        frags = self.bookmarks[name]
        idx   = next((i for i, s in enumerate(frags)
                      if self._fragment_key(s) == key), None)
        if idx is None:
            QMessageBox.information(self, self.t('dlg_bm_info'),
                                    self.t('dlg_not_in_bm'))
            return
        del frags[idx]
        self.search_results = list(frags)
        self.render_results()
        QMessageBox.information(self, self.t('dlg_bm_info'),
                                self.t('dlg_removed_from_bm', name=name))
        self.save_session()

    # ------------------------------------------------------------------
    #  Экспорт подборки
    # ------------------------------------------------------------------
    def export_bookmark(self):
        name = self.bm_combo.currentText()
        if not name:
            QMessageBox.information(self, self.t('dlg_export_title'),
                                    self.t('dlg_no_bm'))
            return
        frags = self.bookmarks.get(name, [])
        if not frags:
            QMessageBox.information(self, self.t('dlg_export_title'),
                                    self.t('dlg_export_empty'))
            return

        path, flt = QFileDialog.getSaveFileName(
            self,
            self.t('dlg_export_save'),
            f"{name}.txt",
            self.t('dlg_export_filters'))
        if not path:
            return

        try:
            ext = os.path.splitext(path)[1].lower()
            heading = self.t('export_heading', name=name)

            if ext == '.txt':
                self._export_txt(path, heading, frags)
            elif ext == '.docx':
                self._export_docx(path, heading, frags)
            elif ext == '.pdf':
                self._export_pdf(path, heading, frags)
            else:
                self._export_txt(path, heading, frags)

            QMessageBox.information(self, self.t('dlg_export_title'),
                                    self.t('dlg_export_done'))
        except Exception as e:
            QMessageBox.critical(self, self.t('dlg_export_title'),
                                 self.t('dlg_export_err', e=e))

    def _export_txt(self, path, heading, frags):
        with open(path, 'w', encoding='utf-8') as f:
            f.write(heading + '\n')
            f.write('=' * 60 + '\n\n')
            for i, d in enumerate(frags, 1):
                book  = d['book']
                start = d.get('start', d.get('index'))
                end   = d.get('end',   d.get('index'))
                text  = '\n'.join(book.paragraphs[start:end + 1])
                f.write(f"[{i}] {book.title}\n{text}\n\n{'—'*40}\n\n")

    def _export_docx(self, path, heading, frags):
        try:
            import docx
        except ImportError:
            raise ImportError("pip install python-docx")
        doc = docx.Document()
        doc.add_heading(heading, 0)
        for i, d in enumerate(frags, 1):
            book  = d['book']
            start = d.get('start', d.get('index'))
            end   = d.get('end',   d.get('index'))
            text  = '\n'.join(book.paragraphs[start:end + 1])
            doc.add_heading(f"{i}. {book.title}", level=2)
            doc.add_paragraph(text)
            doc.add_paragraph('—' * 20)
        doc.save(path)

    def _export_pdf(self, path, heading, frags):
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import (SimpleDocTemplate, Paragraph,
                                            Spacer)
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
        except ImportError:
            raise ImportError("pip install reportlab")

        # Регистрируем шрифт с поддержкой кириллицы
        font_candidates = [
            "DejaVuSans.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "C:/Windows/Fonts/arial.ttf",
        ]
        registered = False
        for fc in font_candidates:
            if os.path.exists(fc):
                pdfmetrics.registerFont(TTFont("MyFont", fc))
                registered = True
                break

        doc    = SimpleDocTemplate(path, pagesize=A4)
        styles = getSampleStyleSheet()
        fn     = "MyFont" if registered else "Helvetica"
        story  = []

        h_style = styles['Heading1']
        h_style.fontName = fn
        p_style = styles['Normal']
        p_style.fontName = fn
        p_style.fontSize = 11
        p_style.leading  = 16

        story.append(Paragraph(heading, h_style))
        story.append(Spacer(1, 12))

        for i, d in enumerate(frags, 1):
            book  = d['book']
            start = d.get('start', d.get('index'))
            end   = d.get('end',   d.get('index'))
            text  = '\n'.join(book.paragraphs[start:end + 1])
            story.append(Paragraph(f"{i}. {book.title}", h_style))
            story.append(Paragraph(
                text.replace('\n', '<br/>'), p_style))
            story.append(Spacer(1, 12))

        doc.build(story)

    # ------------------------------------------------------------------
    #  Сессия
    # ------------------------------------------------------------------
    def save_session(self):
        bm_serial = {}
        for name, frags in self.bookmarks.items():
            bm_serial[name] = [
                {'path':       d.get('path', d['book'].path),
                 'orig_index': d.get('orig_index', d.get('index')),
                 'start':      d.get('start', d.get('index')),
                 'end':        d.get('end',   d.get('index')),
                 'text':       d.get('text', ''),
                 'query':      d.get('query', '')}
                for d in frags]
        data = {
            'lang':      self.lang,
            'books':     [b.path for b in self.books],
            'bookmarks': bm_serial,
        }
        try:
            with open(self.SESSION_FILE, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    def load_session(self):
        if not os.path.exists(self.SESSION_FILE):
            return
        try:
            with open(self.SESSION_FILE, encoding='utf-8') as f:
                data = json.load(f)
        except Exception:
            return

        # Язык
        self.lang = data.get('lang', 'ru')

        # Книги
        for path in data.get('books', []):
            if os.path.exists(path):
                try:
                    b = Book(path)
                    self.books.append(b)
                    self.book_list.addItem(b.title)
                except Exception:
                    pass

        # Подборки
        book_map = {b.path: b for b in self.books}
        for name, frags in data.get('bookmarks', {}).items():
            restored = []
            for s in frags:
                book = book_map.get(s.get('path'))
                if book is None:
                    continue
                restored.append({
                    'book':       book,
                    'path':       s.get('path'),
                    'orig_index': s.get('orig_index'),
                    'index':      s.get('orig_index'),
                    'start':      s.get('start'),
                    'end':        s.get('end'),
                    'text':       s.get('text', ''),
                    'query':      s.get('query', ''),
                })
            self.bookmarks[name] = restored
            self.bm_combo.addItem(name)


# =========================================================================
#  ТОЧКА ВХОДА
# =========================================================================
if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setStyle("Fusion")
    win = Reader()
    win.show()
    sys.exit(app.exec())